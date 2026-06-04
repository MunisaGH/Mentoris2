import os
from django.conf import settings
from groq import Groq
from langchain_groq import ChatGroq
from langchain_community.vectorstores import Chroma
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_core.documents import Document
import logging

logger = logging.getLogger(__name__)

# ══════════════════ [P-10] SINGLETON: Embeddinglar har chaqiruvda qayta yuklanmaslik uchun ══════════════════
_llm_instance = None
_embeddings_instance = None


def get_llm():
    """Groq LLM modelini qaytaradi (singleton)"""
    global _llm_instance
    if _llm_instance is None:
        _llm_instance = ChatGroq(
            groq_api_key=settings.GROQ_API_KEY,
            model_name="llama-3.3-70b-versatile",
            temperature=0.3
        )
    return _llm_instance


def get_embeddings():
    """HuggingFace orqali embedding yaratish (Local & Free, singleton)"""
    global _embeddings_instance
    if _embeddings_instance is None:
        _embeddings_instance = HuggingFaceEmbeddings(model_name="all-MiniLM-L6-v2")
    return _embeddings_instance


import PyPDF2
from docx import Document as DocxDocument


def process_document(file_path, user_id):
    """
    Hujjatni tahlil qilib, vektor bazaga joylash (RAG uchun).
    PDF, DOCX va TXT formatlarini qo'llab-quvvatlaydi.
    """
    text = ""
    try:
        file_path_str = str(file_path).lower()
        if file_path_str.endswith('.pdf'):
            with open(file_path, 'rb') as f:
                pdf_reader = PyPDF2.PdfReader(f)
                for page in pdf_reader.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
        elif file_path_str.endswith('.docx'):
            doc = DocxDocument(file_path)
            text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
        else:
            with open(file_path, 'r', encoding='utf-8') as f:
                text = f.read()
    except Exception as e:
        logger.error(f"Faylni o'qishda xato: {e}")
        return None

    if not text.strip():
        logger.error("Hujjatdan matn ajratib olinmadi.")
        return None

    # Matnni bo'laklarga ajratish
    splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=200)
    chunks = splitter.split_text(text)
    documents = [Document(page_content=chunk, metadata={"user_id": str(user_id)}) for chunk in chunks]

    # Chroma vektor bazaga saqlash
    vectorstore = Chroma.from_documents(
        documents,
        get_embeddings(),
        persist_directory=str(settings.BASE_DIR / "chroma_db")
    )
    return vectorstore
